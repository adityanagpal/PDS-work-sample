#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Mon Nov 12 13:07:37 2018

@author: aditya
"""

from pdf2image import convert_from_path
pages = convert_from_path('c-liat 1227.pdf', 500)
i=0

for page in pages:
    i=i+1
    page.save("c-liat 1227Out"+str(i)+".jpg", 'JPEG')
    
    
     
############# json Implemting ###########################    

import json
from pprint import pprint

with open('jason.json') as f:
    data1 = json.load(f)

with open('jason2.json') as f:
    data2 = json.load(f)
    
with open('jason3.json') as f:
    data3 = json.load(f)    

with open('jason4.json') as f:
    data4 = json.load(f)    

with open('jason5.json') as f:
    data5 = json.load(f)    


x1 = data1["fullTextAnnotation"]["text"]

y1 = data1["textAnnotations"]

a1 = y1[2]
 
if a1['description']=='ORIGINAL':
    print('yes')
else:
   print('No')    
   
for i in range(len(y1)):
    if y1[i]['description'].lower()=='booking':
        x11 = y1[i]['boundingPoly']['vertices']
        break
print('Booking No')
for i in range(16,920):
    if(y1[i]['boundingPoly']['vertices'][0]['x']-x11[0]['x']<20 and y1[i]['boundingPoly']['vertices'][0]['y']-x11[2]['y']<20):
        print(y1[i]['description'])
        break
    
for i in range(len(y1)):
    if y1[i]['description'].lower()=='place':
        x11 = y1[i]['boundingPoly']['vertices']
        break
j=i    
print('Port of Loading')
for i in range(j+1,920):
    if((y1[i]['boundingPoly']['vertices'][0]['x']-x11[0]['x'])<80 and (y1[i]['boundingPoly']['vertices'][0]['y']-x11[2]['y'])<200):
        print(y1[i]['description'])
        break   
    
x2 = data2["fullTextAnnotation"]["text"]

y2 = data2["textAnnotations"]

a2 = y2[2]
 
if a2['description']=='ORIGINAL':
    print('yes')
else:
   print('No')    
  
for i in range(len(y2)):
    if y2[i]['description'].lower()=='b':
        x12 = y2[i]['boundingPoly']['vertices']
        break
print('Booking No')
for i in range(16,920):
    if(y2[i]['boundingPoly']['vertices'][0]['x']-x12[0]['x']<200 and y2[i]['boundingPoly']['vertices'][0]['y']-x12[2]['y']<200):
        print(y2[i]['description'])
        break
    
for i in range(len(y2)):
    if y2[i]['description'].lower()=='port':
        x12 = y2[i]['boundingPoly']['vertices']
        break
j=i    
print('Port of Loading')
for i in range(j+1,len(y2)):
    if((y2[i]['boundingPoly']['vertices'][0]['x']-x12[0]['x'])<300 and (y2[i]['boundingPoly']['vertices'][0]['y']-x12[2]['y'])<200):
        print(y2[i]['description'])
        break       

x3 = data3["fullTextAnnotation"]["text"]

y3 = data3["textAnnotations"]

a3 = y3[2]
 
if a3['description']=='ORIGINAL':
    print('yes')
else:
   print('No')    
'''   
for i in range(len(y3)):
    if y3[i]['description'].lower()=='booking':
        x13 = y3[i]['boundingPoly']['vertices']
        break
print('Booking No')
for i in range(16,len(y3)):
    if(y3[i]['boundingPoly']['vertices'][0]['x']-x13[0]['x']<20 and y3[i]['boundingPoly']['vertices'][0]['y']-x13[2]['y']<20):
        print(y3[i]['description'])
        break
    
for i in range(len(y3)):
    if y3[i]['description'].lower()=='place':
        x13 = y3[i]['boundingPoly']['vertices']
        break
j=i    
print('Port of Loading')
for i in range(j+1,len(y2)):
    if((y2[i]['boundingPoly']['vertices'][0]['x']-x12[0]['x'])<80 and (y2[i]['boundingPoly']['vertices'][0]['y']-x12[2]['y'])<200):
        print(y2[i]['description'])
        break       
'''
x4 = data4["fullTextAnnotation"]["text"]

y4 = data4["textAnnotations"]

a4 = y4[2]
 
if a4['description']=='ORIGINAL':
    print('yes')
else:
   print('No')    

x5 = data5["fullTextAnnotation"]["text"]

y5 = data5["textAnnotations"]

a5 = y5[2]
 
if a4['description']=='ORIGINAL':
    print('yes')
else:
   print('No')    
      
   
from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading("Booking No")
document.add_paragraph("SH8FVF167600")

document.add_heading("Port of Loading")
document.add_paragraph(y1[i]['description'])

document.add_heading("Sea Way Bill No")
document.add_paragraph(y1[23]['description'])

document.add_heading("Sea Way Bill No")
document.add_paragraph(y1[23]['description'])

document.add_heading("Gross Weight")
document.add_paragraph(y1[688]['description'])

document.add_heading("Sea Way Bill No")
document.add_paragraph(y1[23]['description'])

document.add_heading("BL No.")
document.add_paragraph(y2[6]['description'])

document.add_heading("BL No.")
document.add_paragraph(y2[7]['description'])

document.add_heading("BL No.")
document.add_paragraph(y2[8]['description'])

document.add_heading("Date")
document.add_paragraph(y3[5]['description'])

document.add_heading("Date")
document.add_paragraph(y3[6]['description'])

document.add_heading("Date")
document.add_paragraph(y3[7]['description'])

document.add_heading("Date")
document.add_paragraph(y3[8]['description'])

document.add_heading("Date")
document.add_paragraph(y3[9]['description'])

document.add_heading("HBL No")
document.add_paragraph(y3[21]['description'])

i = 22
while i<len(y3):
    document.add_heading(y3[i]['description']+' '+y3[i+1]['description'])
    document.add_paragraph(y3[i+2]['description']+' '+y3[i+3]['description']+' '+y3[i+4]['description']+' '+y3[i+5]['description'])
    i=i+6
    
document.add_heading("Invoice No")
document.add_paragraph(y4[6]['description'])

document.add_heading("Invoice No")
document.add_paragraph(y4[7]['description'])

document.add_heading("Invoice No")
document.add_paragraph(y4[8]['description'])
    
document.add_heading("Packages")
document.add_paragraph(y4[195]['description'])

document.add_heading("Packages")
document.add_paragraph(y4[196]['description'])


    
document.save('new.docx')   
 
doc1 = Document()
doc1.add_paragraph(y1[0]['description'])  
doc2 = Document() 
doc2.add_paragraph(y2[0]['description'])    
doc3 = Document()   
doc3.add_paragraph(y3[0]['description'])    
doc4 = Document()
doc4.add_paragraph(y4[0]['description'])   

doc1.save('new1.docx')    
doc2.save('new2.docx')    
doc3.save('new3.docx')    
doc4.save('new4.docx')








        